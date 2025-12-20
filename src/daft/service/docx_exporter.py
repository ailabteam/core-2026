from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.service.docx_analyzer import (
    ELEMENT_HEADING,
    ELEMENT_LIST,
    ELEMENT_PARAGRAPH,
    ELEMENT_TABLE,
    get_block_text,
    get_block_bbox,
)

logger = logging.getLogger(__name__)

# Alignment mapping
ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

DEFAULT_STYLE_MAP = {
    "heading1": 14,
    "heading2": 13,
    "heading3": 13,
    "body": 13,
    "list": 13,
    "org": 12,
}


def calculate_font_size(
    block: Dict[str, Any],
    style_map: Dict[str, int],
    element_type: str,
    heading_level: int,
) -> int:
    """
    Determine font size based on element type and heading level.
    style_map keys: heading1, heading2, heading3, body, list, org.
    """
    if element_type == ELEMENT_HEADING:
        if heading_level == 1:
            return style_map.get("heading1", DEFAULT_STYLE_MAP["heading1"])
        if heading_level == 2:
            return style_map.get("heading2", DEFAULT_STYLE_MAP["heading2"])
        if heading_level == 3:
            return style_map.get("heading3", DEFAULT_STYLE_MAP["heading3"])
        return style_map.get("heading3", DEFAULT_STYLE_MAP["heading3"])
    if element_type == ELEMENT_LIST:
        return style_map.get("list", DEFAULT_STYLE_MAP["list"])
    return style_map.get("body", DEFAULT_STYLE_MAP["body"])


def add_paragraph_with_format(
    doc: Document,
    text: str,
    element_type: str = ELEMENT_PARAGRAPH,
    heading_level: int = 1,
    alignment: str = "left",
    font_size: int | None = None,
    list_type: str | None = None,
    list_indent: int = 0,
    font_name_body: str = "Times New Roman",
    font_name_heading: str = "Times New Roman",
    font_color: str = "000000",
    bold: bool = False,
) -> None:
    """Add a paragraph to document with appropriate formatting."""
    if not text.strip():
        return

    # Create paragraph
    para = doc.add_paragraph()

    # Set alignment
    if alignment in ALIGNMENT_MAP:
        para.alignment = ALIGNMENT_MAP[alignment]

    # Helper to set run font
    def set_run_font(run):
        run.font.name = font_name_heading if element_type == ELEMENT_HEADING else font_name_body
        run.font.color.rgb = RGBColor.from_string(font_color)
        if font_size:
            run.font.size = Pt(font_size)
        # Apply bold if specified
        run.font.bold = bold

    # Apply formatting based on element type
    if element_type == ELEMENT_HEADING:
        # Use Word heading style
        if heading_level == 1:
            para.style = "Heading 1"
        elif heading_level == 2:
            para.style = "Heading 2"
        elif heading_level == 3:
            para.style = "Heading 3"
        elif heading_level == 4:
            para.style = "Heading 4"
        elif heading_level == 5:
            para.style = "Heading 5"
        else:
            para.style = "Heading 6"

        # Override font size if provided
    elif element_type == ELEMENT_LIST:
        # Apply list style
        if list_type == "bullet":
            para.style = "List Bullet"
        elif list_type == "numbered":
            para.style = "List Number"
        else:
            para.style = "List Bullet"

        # Add indent if needed
        if list_indent > 0:
            para.paragraph_format.left_indent = Inches(list_indent * 0.1)
    else:
        # Regular paragraph
        para.style = "Normal"

    # Add text
    run = para.add_run(text)
    set_run_font(run)


def add_table_from_block(
    doc: Document,
    block: Dict[str, Any],
    font_name_body: str = "Times New Roman",
    font_color: str = "000000",
    font_size: int = 13,
) -> None:
    """Add a table to document from OCR table block."""
    table_data = block.get("table", {})
    rows_data = table_data.get("rows", [])

    if not rows_data:
        return

    # Determine table dimensions
    num_rows = len(rows_data)
    num_cols = max(len(row) for row in rows_data) if rows_data else 0

    if num_rows == 0 or num_cols == 0:
        return

    # Create table
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Light Grid Accent 1"  # Default table style

    # Populate table
    for row_idx, row_data in enumerate(rows_data):
        for col_idx, cell_data in enumerate(row_data):
            if col_idx >= num_cols:
                break
            cell = table.rows[row_idx].cells[col_idx]
            cell_text = cell_data.get("text", "")
            cell.text = cell_text

            # Set cell alignment (default to left) and font
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.name = font_name_body
                run.font.color.rgb = RGBColor.from_string(font_color)
                run.font.size = Pt(font_size)


def export_to_docx(
    pages: List[Dict[str, Any]],
    output_path: Path,
    base_font_size: int = 13,
    detect_columns: bool = True,
    merge_pages: bool = True,
    font_name_body: str = "Times New Roman",
    font_name_heading: str = "Times New Roman",
    font_color: str = "000000",
    style_map: Dict[str, int] | None = None,
) -> None:
    """
    Export OCR results to Word document (.docx).

    Args:
        pages: List of analyzed page dicts (from analyze_page_structure)
        output_path: Output file path
        base_font_size: Base font size in points
        detect_columns: Whether to preserve column layout
        merge_pages: Whether to merge all pages into one document
    """
    doc = Document()

    # Prepare style map
    if style_map is None:
        style_map = DEFAULT_STYLE_MAP
    else:
        # Fill missing keys with defaults
        for k, v in DEFAULT_STYLE_MAP.items():
            style_map.setdefault(k, v)

    # Set default font for Normal style
    normal_style = doc.styles["Normal"]
    normal_style.font.name = font_name_body
    normal_style.font.color.rgb = RGBColor.from_string(font_color)
    normal_style.font.size = Pt(style_map.get("body", base_font_size))
    # Ensure East Asia font set to Times New Roman as well
    normal_style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name_body)

    total_tables = 0
    total_columns = 0

    for page_idx, page in enumerate(pages):
        page_width = page.get("width", 1.0)
        page_height = page.get("height", 1.0)
        blocks = page.get("blocks", [])

        # Add page break if not first page and not merging
        if not merge_pages and page_idx > 0:
            doc.add_page_break()

        # Handle multi-column layout
        has_columns = page.get("has_columns", False)
        column_groups = page.get("columns", [])

        if detect_columns and has_columns and len(column_groups) > 1:
            total_columns += len(column_groups)
            # Create section with columns
            section = doc.sections[-1]
            # Set up columns (Word supports up to 3 columns per section)
            if len(column_groups) == 2:
                # Two-column layout
                sectPr = section._sectPr
                cols = sectPr.xpath("./w:cols")[0] if sectPr.xpath("./w:cols") else None
                if cols is None:
                    cols = sectPr.makeelement(qn("w:cols"))
                    sectPr.append(cols)
                cols.set(qn("w:num"), "2")
                cols.set(qn("w:space"), "720")  # 0.5 inch spacing

            # Process blocks in column order
            # For simplicity, we'll process blocks sequentially
            # Word will handle column flow automatically
            for block in blocks:
                element_type = block.get("element_type", ELEMENT_PARAGRAPH)
                text = get_block_text(block)

                if element_type == ELEMENT_TABLE:
                    add_table_from_block(doc, block)
                    total_tables += 1
                elif text:
                    heading_level = block.get("heading_level", 1)
                    alignment = block.get("alignment", "left")
                    font_size = calculate_font_size(
                        block,
                        style_map=style_map,
                        element_type=element_type,
                        heading_level=heading_level,
                    )
                    list_type = block.get("list_type")
                    list_indent = block.get("list_indent", 0)
                    bold = block.get("bold", False)  # Get bold flag from semantic analysis

                    add_paragraph_with_format(
                        doc,
                        text,
                        element_type=element_type,
                        heading_level=heading_level,
                        alignment=alignment,
                        font_size=font_size,
                        list_type=list_type,
                        list_indent=list_indent,
                        font_name_body=font_name_body,
                        font_name_heading=font_name_heading,
                        font_color=font_color,
                        bold=bold,
                    )
        else:
            # Single column layout
            for block in blocks:
                element_type = block.get("element_type", ELEMENT_PARAGRAPH)
                text = get_block_text(block)

                if element_type == ELEMENT_TABLE:
                    add_table_from_block(doc, block)
                    total_tables += 1
                elif text:
                    heading_level = block.get("heading_level", 1)
                    alignment = block.get("alignment", "left")
                    font_size = calculate_font_size(
                        block,
                        style_map=style_map,
                        element_type=element_type,
                        heading_level=heading_level,
                    )
                    list_type = block.get("list_type")
                    list_indent = block.get("list_indent", 0)
                    bold = block.get("bold", False)  # Get bold flag from semantic analysis

                    add_paragraph_with_format(
                        doc,
                        text,
                        element_type=element_type,
                        heading_level=heading_level,
                        alignment=alignment,
                        font_size=font_size,
                        list_type=list_type,
                        list_indent=list_indent,
                        font_name_body=font_name_body,
                        font_name_heading=font_name_heading,
                        font_color=font_color,
                        bold=bold,
                    )

    # Save document
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    logger.info(
        f"Exported Word document: {output_path} "
        f"(pages: {len(pages)}, tables: {total_tables}, columns detected: {total_columns})"
    )

